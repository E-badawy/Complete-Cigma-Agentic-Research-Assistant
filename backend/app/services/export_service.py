from io import BytesIO

from docx import Document
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph


def export_docx(report: str, references: list[str]):

    document = Document()

    document.add_heading(
        "Research Report",
        level=1
    )

    document.add_paragraph(report)

    document.add_heading(
        "References",
        level=2
    )

    for ref in references:
        document.add_paragraph(
            ref,
            style="List Bullet"
        )

    output = BytesIO()

    document.save(output)

    output.seek(0)

    return output


def export_pdf(report: str, references: list[str]):

    output = BytesIO()

    styles = getSampleStyleSheet()

    pdf = SimpleDocTemplate(output)

    story = []

    story.append(
        Paragraph("<b>Research Report</b>", styles["Heading1"])
    )

    story.append(
        Paragraph(report.replace("\n", "<br/>"), styles["BodyText"])
    )

    story.append(
        Paragraph("<b>References</b>", styles["Heading2"])
    )

    for ref in references:
        story.append(
            Paragraph(ref, styles["BodyText"])
        )

    pdf.build(story)

    output.seek(0)

    return output